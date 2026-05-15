from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "ПРИЛОЖЕНИЕ А. ТЕКСТ ПРОГРАММЫ - основные блоки кода.docx"


FILES: list[tuple[str, str]] = [
    ("manage.py", "Точка входа для выполнения команд Django"),
    ("desktop_app.py", "Запуск desktop-версии приложения"),
    ("config/settings.py", "Настройки Django-проекта"),
    ("config/urls.py", "Основная маршрутизация WEB-приложения"),
    ("config/api_urls.py", "Маршрутизация REST API"),
    ("config/api_serializers.py", "Сериализаторы REST API"),
    ("config/api_viewsets.py", "Представления REST API"),
    ("core/models.py", "Модели учебных групп, дисциплин, отделений, административного журнала"),
    ("core/forms.py", "Формы справочников"),
    ("core/views.py", "Представления главной страницы, справочников, отчетов"),
    ("core/admin_panel.py", "Пользовательская административная панель"),
    ("core/permissions.py", "Проверка прав доступа"),
    ("core/context_processors.py", "Контекстные данные интерфейса"),
    ("core/urls.py", "Маршруты основного приложения"),
    ("users/models.py", "Модели преподавателей, студентов, родителей, диалогов"),
    ("users/forms.py", "Формы пользователей, чатов, импорта студентов"),
    ("users/views.py", "Представления пользователей, личных кабинетов, чатов"),
    ("users/urls.py", "Маршруты приложения пользователей"),
    ("users/management/commands/seed_demo_data.py", "Команда заполнения демонстрационных данных"),
    ("users/management/commands/normalize_teacher_names.py", "Команда нормализации ФИО преподавателей"),
    ("users/management/commands/merge_duplicate_teachers.py", "Команда объединения дублей преподавателей"),
    ("grading/models.py", "Модели курсов, занятий, оценок, посещаемости"),
    ("grading/forms.py", "Формы электронного журнала"),
    ("grading/views.py", "Представления курсов, журнала, посещаемости"),
    ("grading/api_views.py", "API-представления электронного журнала"),
    ("grading/serializers.py", "Сериализаторы оценок и курсов"),
    ("grading/urls.py", "Маршруты приложения электронного журнала"),
    ("reports/forms.py", "Формы параметров отчетов"),
    ("reports/views.py", "Представления формирования отчетов"),
    ("reports/export_utils.py", "Выгрузка отчетов в CSV, Excel, PDF"),
    ("reports/urls.py", "Маршруты отчетов"),
    ("integration/providers/base.py", "Базовый класс поставщика учебных данных"),
    ("integration/providers/mpt_schedule_provider.py", "Получение и парсинг расписания МПТ"),
    ("integration/services/sync_service.py", "Сервис синхронизации справочников"),
    ("integration/management/commands/sync_mpt_catalog.py", "Команда синхронизации справочников"),
    ("integration/management/commands/sync_mpt_groups.py", "Команда синхронизации групп"),
    ("integration/management/commands/sync_mpt_teachers.py", "Команда синхронизации преподавателей"),
    ("templates/base.html", "Базовый HTML-шаблон интерфейса"),
    ("templates/registration/login.html", "Шаблон авторизации"),
    ("core/templates/core/dashboard.html", "Шаблон главной страницы"),
    ("core/templates/core/group_list.html", "Шаблон списка групп"),
    ("core/templates/core/subject_list.html", "Шаблон списка дисциплин"),
    ("core/templates/core/form.html", "Шаблон формы справочника"),
    ("core/templates/core/confirm_delete.html", "Шаблон подтверждения удаления"),
    ("core/templates/core/reports_index.html", "Шаблон страницы отчетов"),
    ("core/templates/core/admin_panel/index.html", "Шаблон главной страницы административной панели"),
    ("core/templates/core/admin_panel/list.html", "Шаблон списков административной панели"),
    ("core/templates/core/admin_panel/form.html", "Шаблон форм административной панели"),
    ("core/templates/core/admin_panel/history.html", "Шаблон истории изменений"),
    ("core/templates/core/admin_panel/confirm_delete.html", "Шаблон удаления в административной панели"),
    ("users/templates/users/student_list.html", "Шаблон списка студентов"),
    ("users/templates/users/student_detail.html", "Шаблон карточки студента"),
    ("users/templates/users/student_course_journal.html", "Шаблон индивидуального журнала студента"),
    ("users/templates/users/teacher_list.html", "Шаблон списка преподавателей"),
    ("users/templates/users/parent_list.html", "Шаблон списка родителей"),
    ("users/templates/users/parent_dashboard.html", "Шаблон кабинета родителя"),
    ("users/templates/users/import_students.html", "Шаблон импорта студентов"),
    ("users/templates/users/chat_list.html", "Шаблон списка диалогов"),
    ("users/templates/users/chat_form.html", "Шаблон формы сообщения"),
    ("users/templates/users/chat_detail.html", "Шаблон страницы диалога"),
    ("grading/templates/grading/course_list.html", "Шаблон списка курсов"),
    ("grading/templates/grading/course_detail.html", "Шаблон страницы курса"),
    ("grading/templates/grading/course_journal.html", "Шаблон электронного журнала"),
    ("grading/templates/grading/attendance_mark.html", "Шаблон отметки посещаемости"),
    ("reports/templates/reports/group_statement.html", "Шаблон ведомости группы"),
    ("reports/templates/reports/top_students.html", "Шаблон рейтинга студентов"),
    ("reports/templates/reports/attendance_report.html", "Шаблон отчета по посещаемости"),
    ("core/tests.py", "Автоматизированные тесты основного приложения"),
    ("users/tests.py", "Автоматизированные тесты пользователей и чатов"),
    ("grading/tests.py", "Автоматизированные тесты электронного журнала"),
    ("integration/tests/test_parsing.py", "Тесты парсинга расписания"),
    ("integration/tests/test_upsert.py", "Тесты синхронизации справочников"),
    ("Dockerfile", "Сборка Docker-образа WEB-приложения"),
    ("docker-compose.yml", "Запуск WEB-приложения и PostgreSQL"),
    ("requirements.txt", "Зависимости локальной версии приложения"),
    ("requirements.docker.txt", "Зависимости Docker-сборки"),
]


def setup_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.0)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(12)

    code_style = styles.add_style("Code Listing", 1)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(7)
    paragraph_format = code_style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_code_line(document: Document, line: str) -> None:
    paragraph = document.add_paragraph(style="Code Listing")
    run = paragraph.add_run(line.replace("\t", "    "))
    run.font.name = "Courier New"
    run.font.size = Pt(7)


def read_text(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8", errors="ignore").splitlines()


def existing_files() -> list[tuple[str, str, int]]:
    result: list[tuple[str, str, int]] = []
    for relative_path, description in FILES:
        path = PROJECT_ROOT / relative_path
        if path.exists():
            result.append((relative_path, description, len(read_text(path))))
    return result


def add_modules_table(document: Document, rows: list[tuple[str, str, int]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["N", "Файл", "Назначение", "Строк"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for number, (relative_path, description, line_count) in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].text = str(number)
        cells[1].text = relative_path
        cells[2].text = description
        cells[3].text = str(line_count)


def add_file_listing(document: Document, number: int, relative_path: str, description: str) -> None:
    path = PROJECT_ROOT / relative_path
    document.add_heading(f"{number}. {relative_path}", level=2)
    add_paragraph(document, description)
    for line in read_text(path):
        add_code_line(document, line)


def main() -> None:
    document = Document()
    setup_document(document)

    document.add_heading("ПРИЛОЖЕНИЕ А. ТЕКСТ ПРОГРАММЫ", level=1)
    add_paragraph(
        document,
        "В приложении приведены основные исходные файлы программного комплекса "
        "«Электронный журнал успеваемости и посещаемости». В состав текста программы "
        "включены серверные модули Django, модели базы данных, формы, представления, "
        "маршруты, REST API, модули отчетности, синхронизации учебных данных, "
        "ключевые HTML-шаблоны, автоматизированные тесты и конфигурационные файлы запуска.",
    )
    add_paragraph(
        document,
        "Служебные миграции, демонстрационные JSON-данные, резервные копии базы данных, "
        "изображения и файлы диаграмм в приложение не включены, так как они не являются "
        "основными блоками программной логики.",
    )

    rows = existing_files()
    document.add_heading("Перечень основных файлов", level=2)
    add_modules_table(document, rows)

    for number, (relative_path, description, _line_count) in enumerate(rows, 1):
        add_file_listing(document, number, relative_path, description)
        if number != len(rows):
            document.add_paragraph("")

    document.save(OUTPUT)
    total_lines = sum(line_count for _path, _description, line_count in rows)
    print(f"Создан файл: {OUTPUT}")
    print(f"Добавлено файлов: {len(rows)}")
    print(f"Строк кода: {total_lines}")


if __name__ == "__main__":
    main()
