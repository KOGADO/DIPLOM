from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = Path.home() / "Downloads" / "Telegram Desktop"
OUTPUT = PROJECT_ROOT / "Приложение_Г_Скрипт_БД_Журнал_МПТ.docx"


ANNOTATION = [
    "В данном разделе представлен скрипт базы данных для программы «Журнал МПТ», который используется для хранения и обработки данных автоматизированной системы электронного журнала успеваемости и посещаемости.",
    "Скрипт базы данных содержит описание структуры таблиц, связей между сущностями, ограничений целостности и индексов, обеспечивающих хранение информации о пользователях, учебных группах, дисциплинах, преподавателях, студентах, родителях, курсах, занятиях, оценках, посещаемости и сообщениях.",
    "Серверная часть приложения реализует взаимодействие с базой данных через модели Django ORM. PostgreSQL применяется в качестве основной системы управления базами данных, обеспечивающей надежное хранение данных, поддержку транзакций, внешних ключей и ограничений целостности.",
]


DICTIONARY_ROWS = [
    ["Ключ", "Наименование поля", "Тип данных", "Ограничение", "Описание"],
    ["1", "2", "3", "4", "5"],
    ["Сущность «auth_user»", "Сущность «auth_user»", "Сущность «auth_user»", "Сущность «auth_user»", "Сущность «auth_user»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID пользователя"],
    ["", "username", "VARCHAR(150)", "NOT NULL, UNIQUE", "Логин пользователя"],
    ["", "password", "VARCHAR(128)", "NOT NULL", "Хешированный пароль пользователя"],
    ["", "first_name", "VARCHAR(150)", "", "Имя пользователя"],
    ["", "last_name", "VARCHAR(150)", "", "Фамилия пользователя"],
    ["", "email", "VARCHAR(254)", "", "Электронная почта"],
    ["", "is_active", "BOOLEAN", "DEFAULT TRUE", "Статус активности пользователя"],
    ["", "is_staff", "BOOLEAN", "DEFAULT FALSE", "Признак доступа к административной части"],
    ["", "is_superuser", "BOOLEAN", "DEFAULT FALSE", "Признак суперпользователя"],
    ["Сущность «core_department»", "Сущность «core_department»", "Сущность «core_department»", "Сущность «core_department»", "Сущность «core_department»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID отделения"],
    ["", "name", "VARCHAR(255)", "NOT NULL, UNIQUE", "Название отделения"],
    ["Сущность «core_group»", "Сущность «core_group»", "Сущность «core_group»", "Сущность «core_group»", "Сущность «core_group»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID учебной группы"],
    ["", "name", "VARCHAR(100)", "NOT NULL, UNIQUE", "Название группы"],
    ["FK", "department_id", "BIGINT", "NULL", "ID отделения"],
    ["FK", "curator_id", "BIGINT", "NULL", "ID куратора группы"],
    ["", "source", "VARCHAR(50)", "DEFAULT 'mpt.ru'", "Источник данных"],
    ["", "external_id", "VARCHAR(255)", "UNIQUE WITH source", "Внешний идентификатор"],
    ["", "is_active", "BOOLEAN", "DEFAULT TRUE", "Статус активности группы"],
    ["Сущность «core_subject»", "Сущность «core_subject»", "Сущность «core_subject»", "Сущность «core_subject»", "Сущность «core_subject»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID дисциплины"],
    ["", "name", "VARCHAR(255)", "NOT NULL, UNIQUE", "Название дисциплины"],
    ["FK", "department_id", "BIGINT", "NULL", "ID отделения"],
    ["", "source", "VARCHAR(50)", "DEFAULT 'mpt.ru'", "Источник данных"],
    ["", "external_id", "VARCHAR(255)", "UNIQUE WITH source", "Внешний идентификатор"],
    ["", "is_active", "BOOLEAN", "DEFAULT TRUE", "Статус активности дисциплины"],
    ["Сущность «users_teacher»", "Сущность «users_teacher»", "Сущность «users_teacher»", "Сущность «users_teacher»", "Сущность «users_teacher»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID преподавателя"],
    ["FK", "user_id", "BIGINT", "NOT NULL, UNIQUE", "ID пользователя"],
    ["FK", "department_id", "BIGINT", "NULL", "ID отделения"],
    ["", "source", "VARCHAR(50)", "DEFAULT 'mpt.ru'", "Источник данных"],
    ["", "external_id", "VARCHAR(255)", "UNIQUE WITH source", "Внешний идентификатор"],
    ["", "is_active", "BOOLEAN", "DEFAULT TRUE", "Статус активности преподавателя"],
    ["Сущность «users_student»", "Сущность «users_student»", "Сущность «users_student»", "Сущность «users_student»", "Сущность «users_student»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID студента"],
    ["FK", "user_id", "BIGINT", "NOT NULL, UNIQUE", "ID пользователя"],
    ["FK", "group_id", "BIGINT", "NOT NULL", "ID учебной группы"],
    ["", "date_of_birth", "DATE", "NULL", "Дата рождения студента"],
    ["Сущность «users_parent»", "Сущность «users_parent»", "Сущность «users_parent»", "Сущность «users_parent»", "Сущность «users_parent»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID родителя"],
    ["FK", "user_id", "BIGINT", "NOT NULL, UNIQUE", "ID пользователя"],
    ["Сущность «users_parent_children»", "Сущность «users_parent_children»", "Сущность «users_parent_children»", "Сущность «users_parent_children»", "Сущность «users_parent_children»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID связи родителя и студента"],
    ["FK", "parent_id", "BIGINT", "NOT NULL", "ID родителя"],
    ["FK", "student_id", "BIGINT", "NOT NULL", "ID студента"],
    ["Сущность «grading_course»", "Сущность «grading_course»", "Сущность «grading_course»", "Сущность «grading_course»", "Сущность «grading_course»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID курса"],
    ["FK", "subject_id", "BIGINT", "NOT NULL", "ID дисциплины"],
    ["FK", "teacher_id", "BIGINT", "NOT NULL", "ID преподавателя"],
    ["FK", "group_id", "BIGINT", "NOT NULL", "ID учебной группы"],
    ["", "semester", "VARCHAR(20)", "NOT NULL", "Семестр"],
    ["", "year", "INTEGER", "NULL", "Учебный год"],
    ["", "source", "VARCHAR(50)", "DEFAULT 'mpt.ru'", "Источник данных"],
    ["", "external_id", "VARCHAR(255)", "UNIQUE WITH source", "Внешний идентификатор"],
    ["", "is_active", "BOOLEAN", "DEFAULT TRUE", "Статус активности курса"],
    ["Сущность «grading_studentcourse»", "Сущность «grading_studentcourse»", "Сущность «grading_studentcourse»", "Сущность «grading_studentcourse»", "Сущность «grading_studentcourse»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID записи на курс"],
    ["FK", "student_id", "BIGINT", "NOT NULL", "ID студента"],
    ["FK", "course_id", "BIGINT", "NOT NULL", "ID курса"],
    ["Сущность «grading_lesson»", "Сущность «grading_lesson»", "Сущность «grading_lesson»", "Сущность «grading_lesson»", "Сущность «grading_lesson»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID занятия"],
    ["FK", "course_id", "BIGINT", "NOT NULL", "ID курса"],
    ["", "date", "DATE", "NOT NULL", "Дата занятия"],
    ["", "topic", "VARCHAR(255)", "DEFAULT ''", "Тема занятия"],
    ["Сущность «grading_lecturetopic»", "Сущность «grading_lecturetopic»", "Сущность «grading_lecturetopic»", "Сущность «grading_lecturetopic»", "Сущность «grading_lecturetopic»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID темы лекции"],
    ["FK", "course_id", "BIGINT", "NOT NULL", "ID курса"],
    ["", "title", "VARCHAR(255)", "NOT NULL", "Название темы"],
    ["", "order", "INTEGER", "DEFAULT 0", "Порядковый номер"],
    ["", "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Дата создания"],
    ["Сущность «grading_attendance»", "Сущность «grading_attendance»", "Сущность «grading_attendance»", "Сущность «grading_attendance»", "Сущность «grading_attendance»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID посещаемости"],
    ["FK", "lesson_id", "BIGINT", "NOT NULL", "ID занятия"],
    ["FK", "student_id", "BIGINT", "NOT NULL", "ID студента"],
    ["", "status", "VARCHAR(10)", "NOT NULL", "Статус посещаемости"],
    ["", "comment", "VARCHAR(255)", "DEFAULT ''", "Комментарий"],
    ["Сущность «grading_grade»", "Сущность «grading_grade»", "Сущность «grading_grade»", "Сущность «grading_grade»", "Сущность «grading_grade»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID оценки"],
    ["FK", "student_id", "BIGINT", "NOT NULL", "ID студента"],
    ["FK", "course_id", "BIGINT", "NOT NULL", "ID курса"],
    ["", "grade_type", "VARCHAR(10)", "DEFAULT 'grade'", "Тип оценки"],
    ["", "value", "SMALLINT", "CHECK 2-5", "Значение оценки"],
    ["", "date", "DATE", "NOT NULL", "Дата выставления оценки"],
    ["", "comment", "VARCHAR(255)", "DEFAULT ''", "Комментарий"],
    ["Сущность «users_chatdialog»", "Сущность «users_chatdialog»", "Сущность «users_chatdialog»", "Сущность «users_chatdialog»", "Сущность «users_chatdialog»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID диалога"],
    ["FK", "student_id", "BIGINT", "NOT NULL", "ID студента"],
    ["FK", "teacher_id", "BIGINT", "NOT NULL", "ID преподавателя"],
    ["FK", "related_grade_id", "BIGINT", "NULL", "ID оценки, к которой относится вопрос"],
    ["", "title", "VARCHAR(255)", "NOT NULL", "Заголовок диалога"],
    ["", "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Дата создания"],
    ["", "updated_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Дата обновления"],
    ["Сущность «users_chatmessage»", "Сущность «users_chatmessage»", "Сущность «users_chatmessage»", "Сущность «users_chatmessage»", "Сущность «users_chatmessage»"],
    ["PK", "id", "BIGINT", "NOT NULL", "ID сообщения"],
    ["FK", "chat_id", "BIGINT", "NOT NULL", "ID диалога"],
    ["FK", "sender_id", "BIGINT", "NOT NULL", "ID отправителя"],
    ["", "sender_role", "VARCHAR(16)", "NOT NULL", "Роль отправителя"],
    ["", "message", "TEXT", "NOT NULL", "Текст сообщения"],
    ["", "attachment", "VARCHAR(100)", "NULL", "Прикрепленный файл"],
    ["", "created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Дата создания"],
    ["", "is_read", "BOOLEAN", "DEFAULT FALSE", "Статус прочтения"],
]


SQL_LINES = """
-- База данных автоматизированной системы электронного журнала успеваемости и посещаемости
-- СУБД: PostgreSQL

CREATE TABLE auth_user (
    id BIGSERIAL PRIMARY KEY,
    password VARCHAR(128) NOT NULL,
    last_login TIMESTAMP NULL,
    is_superuser BOOLEAN NOT NULL DEFAULT FALSE,
    username VARCHAR(150) NOT NULL UNIQUE,
    first_name VARCHAR(150) NOT NULL DEFAULT '',
    last_name VARCHAR(150) NOT NULL DEFAULT '',
    email VARCHAR(254) NOT NULL DEFAULT '',
    is_staff BOOLEAN NOT NULL DEFAULT FALSE,
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    date_joined TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE auth_group (
    id BIGSERIAL PRIMARY KEY,
    name VARCHAR(150) NOT NULL UNIQUE
);

CREATE TABLE auth_user_groups (
    id BIGSERIAL PRIMARY KEY,
    user_id BIGINT NOT NULL REFERENCES auth_user(id) ON DELETE CASCADE,
    group_id BIGINT NOT NULL REFERENCES auth_group(id) ON DELETE CASCADE,
    CONSTRAINT uq_auth_user_groups UNIQUE (user_id, group_id)
);

CREATE TABLE core_department (
    id BIGSERIAL PRIMARY KEY,
    name VARCHAR(255) NOT NULL UNIQUE
);

CREATE TABLE core_group (
    id BIGSERIAL PRIMARY KEY,
    name VARCHAR(100) NOT NULL UNIQUE,
    department_id BIGINT NULL REFERENCES core_department(id) ON DELETE SET NULL,
    curator_id BIGINT NULL REFERENCES auth_user(id) ON DELETE SET NULL,
    source VARCHAR(50) NOT NULL DEFAULT 'mpt.ru',
    external_id VARCHAR(255) NULL,
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    CONSTRAINT uq_core_group_source_external UNIQUE (source, external_id)
);

CREATE TABLE core_subject (
    id BIGSERIAL PRIMARY KEY,
    name VARCHAR(255) NOT NULL UNIQUE,
    department_id BIGINT NULL REFERENCES core_department(id) ON DELETE SET NULL,
    source VARCHAR(50) NOT NULL DEFAULT 'mpt.ru',
    external_id VARCHAR(255) NULL,
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    CONSTRAINT uq_core_subject_source_external UNIQUE (source, external_id)
);

CREATE TABLE users_teacher (
    id BIGSERIAL PRIMARY KEY,
    user_id BIGINT NOT NULL UNIQUE REFERENCES auth_user(id) ON DELETE CASCADE,
    department_id BIGINT NULL REFERENCES core_department(id) ON DELETE SET NULL,
    source VARCHAR(50) NOT NULL DEFAULT 'mpt.ru',
    external_id VARCHAR(255) NULL,
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    CONSTRAINT uq_teacher_source_external UNIQUE (source, external_id)
);

CREATE TABLE users_student (
    id BIGSERIAL PRIMARY KEY,
    user_id BIGINT NOT NULL UNIQUE REFERENCES auth_user(id) ON DELETE CASCADE,
    group_id BIGINT NOT NULL REFERENCES core_group(id) ON DELETE CASCADE,
    date_of_birth DATE NULL
);

CREATE TABLE users_parent (
    id BIGSERIAL PRIMARY KEY,
    user_id BIGINT NOT NULL UNIQUE REFERENCES auth_user(id) ON DELETE CASCADE
);

CREATE TABLE users_parent_children (
    id BIGSERIAL PRIMARY KEY,
    parent_id BIGINT NOT NULL REFERENCES users_parent(id) ON DELETE CASCADE,
    student_id BIGINT NOT NULL REFERENCES users_student(id) ON DELETE CASCADE,
    CONSTRAINT uq_parent_student UNIQUE (parent_id, student_id)
);

CREATE TABLE grading_course (
    id BIGSERIAL PRIMARY KEY,
    subject_id BIGINT NOT NULL REFERENCES core_subject(id) ON DELETE CASCADE,
    teacher_id BIGINT NOT NULL REFERENCES users_teacher(id) ON DELETE CASCADE,
    group_id BIGINT NOT NULL REFERENCES core_group(id) ON DELETE CASCADE,
    semester VARCHAR(20) NOT NULL,
    year INTEGER NULL,
    source VARCHAR(50) NOT NULL DEFAULT 'mpt.ru',
    external_id VARCHAR(255) NULL,
    is_active BOOLEAN NOT NULL DEFAULT TRUE,
    CONSTRAINT uq_course_main UNIQUE (subject_id, teacher_id, group_id, semester),
    CONSTRAINT uq_course_source_external UNIQUE (source, external_id)
);

CREATE TABLE grading_studentcourse (
    id BIGSERIAL PRIMARY KEY,
    student_id BIGINT NOT NULL REFERENCES users_student(id) ON DELETE CASCADE,
    course_id BIGINT NOT NULL REFERENCES grading_course(id) ON DELETE CASCADE,
    CONSTRAINT uq_student_course UNIQUE (student_id, course_id)
);

CREATE TABLE grading_lesson (
    id BIGSERIAL PRIMARY KEY,
    course_id BIGINT NOT NULL REFERENCES grading_course(id) ON DELETE CASCADE,
    date DATE NOT NULL,
    topic VARCHAR(255) NOT NULL DEFAULT ''
);

CREATE TABLE grading_lecturetopic (
    id BIGSERIAL PRIMARY KEY,
    course_id BIGINT NOT NULL REFERENCES grading_course(id) ON DELETE CASCADE,
    title VARCHAR(255) NOT NULL,
    "order" INTEGER NOT NULL DEFAULT 0,
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    CONSTRAINT uq_lecturetopic_course_title UNIQUE (course_id, title)
);

CREATE TABLE grading_attendance (
    id BIGSERIAL PRIMARY KEY,
    lesson_id BIGINT NOT NULL REFERENCES grading_lesson(id) ON DELETE CASCADE,
    student_id BIGINT NOT NULL REFERENCES users_student(id) ON DELETE CASCADE,
    status VARCHAR(10) NOT NULL,
    comment VARCHAR(255) NOT NULL DEFAULT '',
    CONSTRAINT uq_attendance_lesson_student UNIQUE (lesson_id, student_id)
);

CREATE TABLE grading_grade (
    id BIGSERIAL PRIMARY KEY,
    student_id BIGINT NOT NULL REFERENCES users_student(id) ON DELETE CASCADE,
    course_id BIGINT NOT NULL REFERENCES grading_course(id) ON DELETE CASCADE,
    grade_type VARCHAR(10) NOT NULL DEFAULT 'grade',
    value SMALLINT NOT NULL CHECK (value BETWEEN 2 AND 5),
    date DATE NOT NULL,
    comment VARCHAR(255) NOT NULL DEFAULT ''
);

CREATE TABLE users_chatdialog (
    id BIGSERIAL PRIMARY KEY,
    student_id BIGINT NOT NULL REFERENCES users_student(id) ON DELETE CASCADE,
    teacher_id BIGINT NOT NULL REFERENCES users_teacher(id) ON DELETE CASCADE,
    related_grade_id BIGINT NULL REFERENCES grading_grade(id) ON DELETE SET NULL,
    title VARCHAR(255) NOT NULL,
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE users_chatmessage (
    id BIGSERIAL PRIMARY KEY,
    chat_id BIGINT NOT NULL REFERENCES users_chatdialog(id) ON DELETE CASCADE,
    sender_id BIGINT NOT NULL REFERENCES auth_user(id) ON DELETE CASCADE,
    sender_role VARCHAR(16) NOT NULL,
    message TEXT NOT NULL,
    attachment VARCHAR(100) NULL,
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
    is_read BOOLEAN NOT NULL DEFAULT FALSE
);

CREATE TABLE core_adminlog (
    id BIGSERIAL PRIMARY KEY,
    user_id BIGINT NULL REFERENCES auth_user(id) ON DELETE SET NULL,
    content_type_id INTEGER NULL,
    object_id VARCHAR(64) NOT NULL,
    object_repr VARCHAR(255) NOT NULL,
    action VARCHAR(12) NOT NULL,
    changed_fields JSONB NOT NULL DEFAULT '{}'::jsonb,
    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
);

CREATE INDEX idx_core_group_external_id ON core_group(external_id);
CREATE INDEX idx_core_subject_external_id ON core_subject(external_id);
CREATE INDEX idx_teacher_external_id ON users_teacher(external_id);
CREATE INDEX idx_course_external_id ON grading_course(external_id);
CREATE INDEX idx_grade_student_course ON grading_grade(student_id, course_id);
CREATE INDEX idx_lesson_course_date ON grading_lesson(course_id, date);
CREATE INDEX idx_attendance_student ON grading_attendance(student_id);
CREATE INDEX idx_chatdialog_student_teacher ON users_chatdialog(student_id, teacher_id);
CREATE INDEX idx_chatmessage_chat_created ON users_chatmessage(chat_id, created_at);

INSERT INTO auth_group (name) VALUES
    ('Admin'),
    ('Teacher'),
    ('Teachers'),
    ('Student'),
    ('Parent')
ON CONFLICT (name) DO NOTHING;
""".strip().splitlines()


def find_source() -> Path:
    target_name = "ПРИЛОЖЕНИЕ_Г. СКРИПТ БД.docx"
    exact = SOURCE_DIR / target_name
    if exact.exists():
        return exact
    matches = [p for p in SOURCE_DIR.glob("*.docx") if "СКРИПТ БД" in p.name]
    if not matches:
        raise FileNotFoundError("Не найден исходный документ со скриптом БД")
    return matches[0]


def set_paragraph(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def set_cell(cell, text: str) -> None:
    cell.text = text


def set_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        row = table.rows[-1]
        row._tr.getparent().remove(row._tr)
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            set_cell(table.rows[row_index].cells[col_index], value)
        if row_index >= 2 and row_data[0].startswith("Сущность «") and len(set(row_data)) == 1:
            merged = table.rows[row_index].cells[0].merge(table.rows[row_index].cells[-1])
            merged.text = row_data[0]
            for paragraph in merged.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def style_script_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def main() -> None:
    doc = Document(find_source())

    set_paragraph(doc.paragraphs[0], "ПРИЛОЖЕНИЕ Г. СКРИПТ БАЗЫ ДАННЫХ")
    for i, text in enumerate(ANNOTATION, start=2):
        set_paragraph(doc.paragraphs[i], text)
    set_paragraph(doc.paragraphs[9], "Наименование «Журнал МПТ».")
    set_paragraph(
        doc.paragraphs[11],
        "Программа предназначена для доступа к автоматизированной системе электронного журнала успеваемости и посещаемости. Система обеспечивает авторизацию пользователей, разграничение доступа по ролям, управление учебными группами, дисциплинами, преподавателями, студентами, родителями, курсами, занятиями, оценками, посещаемостью, сообщениями и отчетами.",
    )
    set_paragraph(doc.paragraphs[14], "Таблица 1 – Словарь данных")

    set_table(doc.tables[0], DICTIONARY_ROWS)

    for paragraph in list(doc.paragraphs[16:]):
        remove_paragraph(paragraph)

    for line in SQL_LINES:
        paragraph = doc.add_paragraph(line)
        style_script_paragraph(paragraph)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
